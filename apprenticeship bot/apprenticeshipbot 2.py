from docx import Document

name = ""
street = ""
city = ""

doc = Document()

#informations
norm = doc.add_paragraph("Kilian Niebling")
norm.add_run().add_break()
norm.add_run("Pfarrhausstrasse 20")
norm.add_run().add_break()
norm.add_run("8344 Bäretswil")
norm.add_run().add_break()
norm.add_run("Tel. 076 766 09 15")

norm = doc.add_paragraph(name)
norm.add_run().add_break()
norm.add_run(street)
norm.add_run().add_break()
norm.add_run(city)

norm.style.font.name = "Arial"
norm.paragraph_format.line_spacing = 1


#title
head = doc.add_heading(level=1)
head_run = head.add_run('Bewerbung für die Lehrstelle als Informatiker (Applikationsentwickler)')
head_run.underline = True


doc.save("test.docx")